"""Build v2 docx by reusing v1 template (preserve official styles).

Strategy:
1. Open v1 docx as template (keeps fonts, heading sizes, table styles, page margins)
2. Wipe the body (paragraphs + tables) but preserve sectPr (page setup)
3. Re-insert content from markdown draft using v1's built-in Heading styles
4. Images use v1's standard image style; tables use v1's Table1 style

No identifier color bands, no hero images, no fancy typography.
"""
from __future__ import annotations
import os, re, sys, shutil, copy
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


REPO = Path("/Volumes/Tim aaddtional/Download/115資通訊/tedt")
COMP = REPO / "docs" / "competition"
DRAFTS = COMP / "drafts"
PLACEHOLDERS = COMP / "assets" / "placeholders"
MERMAID = COMP / "assets" / "mermaid"
MERMAID_DYN = COMP / "assets" / "mermaid-dynamic"
NOTION_PHOTOS = COMP / "assets" / "notion-photos"
AI_GEN = COMP / "assets" / "ai-generated"
SCREENSHOTS_PAGES = REPO / "pages-dist" / "screenshots"
ASSETS_SCREENSHOTS = REPO / "assets" / "screenshots"
SOURCE_V1 = Path("/Volumes/Tim aaddtional/Download/準備比賽")
OUT = Path("/Volumes/Tim aaddtional/Download/準備比賽")

APP_CONFIG = {
    "app1": {
        "src": DRAFTS / "app1-rewrite-draft.md",
        "template": SOURCE_V1 / "初版作品說明書(國) (1).docx",
        "out": OUT / "初版作品說明書(國)_v2.docx",
        "level": "國小",
        "name": "AI 智慧型白板機器人",
        "keywords": "LLM、語音辨識、手寫辨識與姿態辨識",
    },
    "app2": {
        "src": DRAFTS / "app2-rewrite-draft.md",
        "template": SOURCE_V1 / "初版作品說明書(品) (1).docx",
        "out": OUT / "初版作品說明書(品)_v2.docx",
        "level": "國小",
        "name": "校園 AI 多功能服務機器人",
        "keywords": "配送系統、AI 技術、避障技術",
    },
    "app3": {
        "src": DRAFTS / "app3-rewrite-draft.md",
        "template": SOURCE_V1 / "初版作品說明書(印).docx",
        "out": OUT / "初版作品說明書(印)_v2.docx",
        "level": "國中",
        "name": "AI 校園心靈守護者",
        "keywords": "AI 情緒辨識、智慧校園、預警系統",
    },
}

IMG_RE = re.compile(r"^\s*\[IMG:\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*(.+?)\s*\]\s*$")
TABLE_HEAD = re.compile(r"^\s*\|.*\|\s*$")
TABLE_SEP = re.compile(r"^\s*\|(?:\s*:?-+:?\s*\|)+\s*$")
LIST_RE = re.compile(r"^(\s*)([-*]|\d+\.)\s+(.+)$")
BLOCK_QUOTE = re.compile(r"^>\s?(.*)$")
CODE_FENCE = re.compile(r"^```(.*)$")
HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")

FONT_LATIN = "Arial"
# Use Word's standard Traditional Chinese fonts. PMingLiU = 新細明體 (Word default,
# bundled with both Windows Office and Mac Office). Avoid PingFang TC because
# LibreOffice on macOS falls back to a handwriting face when it doesn't know it.
FONT_CJK = "PMingLiU"
FONT_CODE = "Menlo"
BODY_SIZE = Pt(11)
CAPTION_SIZE = Pt(9)
LINE_SPACING = 1.15
TABLE_FILL = "F2F2F2"
BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(64, 64, 64)


def set_run_font(run, size=None, bold=None, italic=None, color=BLACK, code=False):
    """Apply deterministic font tokens; do not rely on Word theme defaults."""
    run.font.name = FONT_CODE if code else FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_CODE if code else FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_CODE if code else FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CJK)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_rhythm(p, before=Pt(0), after=Pt(6), line_spacing=LINE_SPACING):
    pf = p.paragraph_format
    pf.space_before = before
    pf.space_after = after
    pf.line_spacing = line_spacing
    pf.widow_control = True


def configure_styles(doc):
    """Normalize the reused v1 template into the final official style contract."""
    styles = doc.styles

    def set_style_font(style_name, size, bold=None, color=BLACK, before=None, after=None):
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style.font.size = size
        style.font.bold = bold
        style.font.color.rgb = color
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), FONT_LATIN)
        rfonts.set(qn("w:hAnsi"), FONT_LATIN)
        rfonts.set(qn("w:eastAsia"), FONT_CJK)
        pf = style.paragraph_format
        pf.line_spacing = LINE_SPACING
        if before is not None:
            pf.space_before = before
        if after is not None:
            pf.space_after = after

    set_style_font("Normal", BODY_SIZE, bold=False, color=BLACK, before=Pt(0), after=Pt(6))
    set_style_font("Heading 1", Pt(16), bold=False, color=BLACK, before=Pt(12), after=Pt(4))
    set_style_font("Heading 2", Pt(13), bold=False, color=BLACK, before=Pt(8), after=Pt(3))
    set_style_font("Heading 3", Pt(12), bold=False, color=DARK_GRAY, before=Pt(6), after=Pt(3))
    try:
        set_style_font("Heading 4", Pt(11), bold=False, color=DARK_GRAY, before=Pt(6), after=Pt(3))
    except KeyError:
        pass

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def _paragraph_contains_drawing(p):
    return bool(p._p.xpath(".//w:drawing"))


def _set_cover_run_fonts(p, size, bold=False):
    for r in p.runs:
        set_run_font(r, size=size, bold=bold, color=BLACK)


def normalize_doc_fonts(doc):
    """Final pass over generated content, including direct-format runs."""
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        pf = p.paragraph_format
        pf.widow_control = True

        if idx == 0 and not text:
            set_paragraph_rhythm(p, after=Pt(42))
            continue

        if _paragraph_contains_drawing(p):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_rhythm(p, before=Pt(6), after=Pt(2))
            p.paragraph_format.keep_with_next = True
            continue

        if text.startswith("圖："):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_rhythm(p, before=Pt(0), after=Pt(8))
            for r in p.runs:
                set_run_font(r, size=CAPTION_SIZE, italic=True, color=DARK_GRAY)
            continue

        if text in {"臺北市 114 年度中小學資通訊應用大賽", "智組型機器人", "創意賽作品說明書", "智組型機器人創意賽作品說明書"}:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_after = Pt(26) if text in {"創意賽作品說明書", "智組型機器人創意賽作品說明書"} else Pt(2 if text.startswith("臺北市") else 0)
            set_paragraph_rhythm(p, before=Pt(0), after=title_after)
            _set_cover_run_fonts(p, Pt(20), bold=False)
            continue

        if text.startswith("學 層 別：") or text.startswith("作品名稱："):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_rhythm(p, before=Pt(10), after=Pt(0), line_spacing=1.08)
            _set_cover_run_fonts(p, Pt(16), bold=False)
            continue

        if text.startswith("關 鍵 詞：") or text.startswith("編") and "號：" in text:
            set_paragraph_rhythm(p, before=Pt(22 if text.startswith("關") else 4), after=Pt(0), line_spacing=1.08)
            p.paragraph_format.left_indent = Cm(0.65)
            _set_cover_run_fonts(p, Pt(14 if text.startswith("關") else 16), bold=False)
            continue

        if text == "製作說明：":
            set_paragraph_rhythm(p, before=Pt(24), after=Pt(8))
            _set_cover_run_fonts(p, Pt(12), bold=False)
            continue

        if text in {
            "1. 說明書封面僅寫組別、年級別、作品名稱及關鍵詞。",
            "2. 編號由承辦單位統一編列。",
            "3. 封面由選手自行設計。",
        }:
            set_paragraph_rhythm(p, before=Pt(6), after=Pt(0))
            p.paragraph_format.left_indent = Cm(1.9)
            p.paragraph_format.first_line_indent = Cm(-0.32)
            _set_cover_run_fonts(p, Pt(12), bold=False)
            continue

        if p.style.name.startswith("Heading"):
            level = int(p.style.name.rsplit(" ", 1)[-1]) if p.style.name[-1].isdigit() else 1
            size = {1: Pt(16), 2: Pt(13), 3: Pt(12)}.get(level, Pt(11))
            color = DARK_GRAY if level >= 3 else BLACK
            before = {1: Pt(12), 2: Pt(8), 3: Pt(6)}.get(level, Pt(6))
            after = {1: Pt(4), 2: Pt(3), 3: Pt(3)}.get(level, Pt(3))
            set_paragraph_rhythm(p, before=before, after=after)
            p.paragraph_format.keep_with_next = True
            for r in p.runs:
                set_run_font(r, size=size, bold=False, color=color)
        elif p.paragraph_format.first_line_indent is not None and p.paragraph_format.first_line_indent.cm < 0:
            set_paragraph_rhythm(p, after=Pt(2), line_spacing=1.12)
            p.paragraph_format.keep_together = True
            for r in p.runs:
                is_code = r.font.name == FONT_CODE
                set_run_font(r, size=Pt(10) if is_code else BODY_SIZE, code=is_code)
        else:
            set_paragraph_rhythm(p)
            for r in p.runs:
                # Preserve code runs that were explicitly set to Menlo.
                is_code = r.font.name == FONT_CODE
                set_run_font(r, size=Pt(10) if is_code else BODY_SIZE, code=is_code)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for p in cell.paragraphs:
                    set_paragraph_rhythm(p, after=Pt(2), line_spacing=1.1)
                    for r in p.runs:
                        set_run_font(r, size=r.font.size or Pt(10), bold=r.bold, italic=r.italic, color=BLACK)


def wipe_body(doc):
    """Remove all paragraphs and tables from body, keep sectPr (page setup)."""
    body = doc.element.body
    # collect children to remove (everything except sectPr)
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def available_width_twips(doc):
    section = doc.sections[0]
    # python-docx Length subtraction returns EMUs; 1 twip = 635 EMUs.
    return int((section.page_width - section.left_margin - section.right_margin) / 635)


def resolve_image(spec, app_key, line_no=None):
    name, kind, ref = spec
    kind = kind.strip().lower()
    ref = ref.strip()

    # PRIORITY: for mermaid, check inline-rendered first (precise per-line match)
    if kind == "mermaid" and line_no is not None:
        cand = MERMAID_DYN / f"{app_key}-inline-L{line_no:04d}.png"
        if cand.exists():
            return cand

    if kind == "real_screenshot":
        parts = re.split(r"[、,]| 或 | or ", ref)
        for raw in parts:
            raw = raw.strip().strip(";；。.").strip()
            if not raw:
                continue
            for fn in re.findall(r"[\w\-/.]+\.(?:png|jpg|jpeg|svg)", raw, re.I):
                base = Path(fn).name
                for root in [REPO, SCREENSHOTS_PAGES, ASSETS_SCREENSHOTS,
                             REPO / "pages-dist", REPO / "pages-dist" / "screenshots"]:
                    cand = root / base
                    if cand.exists():
                        return cand
                for cand in REPO.glob(base):
                    if cand.is_file():
                        return cand
                for cand in (REPO / "pages-dist").rglob(base):
                    return cand
        return None

    if kind == "mermaid":
        # Chinese keyword → english key mapping
        target = name + " " + ref
        zh_to_key = [
            (["狀態圖", "狀態", "生命週期"], "state"),
            (["操作流程", "流程圖"], "flow"),
            (["校園地圖", "任務地圖", "校園任務", "地圖"], "map"),
            (["時間軸", "歷程"], "timeline"),
            (["三層架構", "系統架構", "架構圖", "硬體與軟體"], "arch"),
        ]
        key = None
        for zh_list, ek in zh_to_key:
            if any(z in target for z in zh_list):
                key = ek; break
        if key:
            # exact stem suffix match
            for f in MERMAID.glob(f"{app_key}-{key}.png"):
                if f.exists():
                    return f
        # try English suffix
        target_l = target.lower()
        for f in MERMAID.glob(f"{app_key}-*.png"):
            stem_key = f.stem.split("-")[-1]
            if stem_key in target_l:
                return f
        # do NOT fallback to first file - return None to render as missing note
        return None

    if kind == "ai_concept":
        # Per user feedback: NO color banners. Always render as text note.
        return None

    if kind == "placeholder_real":
        target = (name + " " + ref).lower()
        keywords = {
            "machine-front": ["主視角", "正面", "實機", "全機", "外觀"],
            "machine-wiring": ["接線", "電路", "佈線"],
            "demo-live": ["demo", "現場", "示範", "擦拭"],
            "delivery-loop": ["配送"],
            "fleet-multi": ["多機", "車隊", "4 台"],
            "sensor-node": ["感測節點", "節點"],
            "patrol-bot": ["巡邏"],
            "deploy-scene": ["部署", "場域", "走廊", "教室"],
        }
        scoring = []
        for f in PLACEHOLDERS.glob(f"{app_key}-placeholder-*.png"):
            key = f.stem.split("-", 2)[-1]
            score = 0
            for kw in keywords.get(key, [key]):
                if kw in target:
                    score += 2
            scoring.append((score, f))
        scoring.sort(key=lambda x: -x[0])
        if scoring:
            return scoring[0][1]
        return None

    if kind == "notion_photo":
        for f in NOTION_PHOTOS.glob(f"{app_key}-*"):
            return f
        return None
    return None


def fit_image_size_cm(path):
    """Compute width in cm that fits page (max 14cm wide, max 22cm tall)."""
    from PIL import Image
    try:
        img = Image.open(str(path))
        w_px, h_px = img.size
    except Exception:
        return 14.0  # default fallback
    if h_px == 0:
        return 14.0
    ratio = w_px / h_px  # >1 = wide, <1 = tall
    MAX_W = 14.2  # cm, within A4 / 2.54cm margins
    MAX_H = 20.5  # cm, leave room for caption and surrounding text
    # if image is wider than the page aspect ratio (14/22 = 0.636), width-limited
    if ratio >= MAX_W / MAX_H:
        return MAX_W
    else:
        # height-limited
        return round(MAX_H * ratio, 2)


def add_image_with_caption(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    try:
        w_cm = fit_image_size_cm(path)
        run.add_picture(str(path), width=Cm(w_cm))
    except Exception:
        return False
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"圖：{caption}")
    set_run_font(cr, size=CAPTION_SIZE, italic=True, color=DARK_GRAY)
    cap.paragraph_format.space_after = Pt(8)
    return True


def add_missing_image_note(doc, name, kind, ref):
    """A simple gray-bordered text box (no color)."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    _set_table_layout_fixed(t, available_width_twips(doc))
    cell = t.rows[0].cells[0]
    _set_row_cant_split(t.rows[0])
    _set_cell_margins(cell)
    # Add border via XML
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "dashed")
        b.set(qn("w:sz"), "8")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), "808080")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    # gray fill
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), "F5F5F5")
    tcPr.append(shd)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"〔此處請補上：{name}〕")
    set_run_font(r, size=Pt(10), bold=True, color=DARK_GRAY)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    suggestion = ref if len(ref) < 80 else ref[:80] + "…"
    r2 = p2.add_run(f"建議來源：{suggestion}")
    set_run_font(r2, size=Pt(9), italic=True, color=DARK_GRAY)


def add_image(doc, name, kind, ref, app_key, line_no=None):
    path = resolve_image((name, kind, ref), app_key, line_no=line_no)
    if path and path.exists():
        if add_image_with_caption(doc, path, name):
            return
    add_missing_image_note(doc, name, kind, ref)


def parse_inline_into_paragraph(p, text):
    cur = 0
    for m in re.finditer(r"\*\*([^*]+)\*\*|`([^`]+)`", text):
        if m.start() > cur:
            r = p.add_run(text[cur:m.start()])
            set_run_font(r, size=BODY_SIZE)
        if m.group(1):
            r = p.add_run(m.group(1))
            set_run_font(r, size=BODY_SIZE, bold=True)
        else:
            r = p.add_run(m.group(2))
            set_run_font(r, size=Pt(10), code=True)
        cur = m.end()
    if cur < len(text):
        r = p.add_run(text[cur:])
        set_run_font(r, size=BODY_SIZE)


def _set_cell_border(cell, color="000000", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_table_layout_fixed(table, total_twips):
    """Lock the table to fixed layout so cell widths actually take effect."""
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    # set tblLayout type=fixed
    existing = tblPr.find(qn("w:tblLayout"))
    if existing is not None:
        tblPr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    # set tblW (total table width)
    existing_w = tblPr.find(qn("w:tblW"))
    if existing_w is not None:
        tblPr.remove(existing_w)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _set_table_grid(table, widths):
    existing = table._element.find(qn("w:tblGrid"))
    if existing is not None:
        table._element.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tblPr = table._element.find(qn("w:tblPr"))
    insert_at = 1 if tblPr is not None else 0
    table._element.insert(insert_at, grid)


def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    margins = {
        "top": top,
        "start": start,
        "left": start,
        "bottom": bottom,
        "end": end,
        "right": end,
    }
    for edge, val in margins.items():
        node = tcMar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tcMar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        trPr.append(tbl_header)


def _set_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def display_len(text):
    return sum(2 if ord(ch) > 127 else 1 for ch in text)


def add_md_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    # max content length per col → use to allocate widths
    max_lens = [0] * cols
    for row in rows:
        for ci in range(cols):
            if ci < len(row):
                max_lens[ci] = max(max_lens[ci], display_len(row[ci].strip()))
    total = sum(max_lens) or cols
    AVAIL_TWIPS = max(7200, available_width_twips(doc) - 120)
    min_width = 760 if cols >= 4 else 1100
    col_twips = [max(min_width, int(AVAIL_TWIPS * (l / total))) for l in max_lens]
    # re-normalize so sum equals AVAIL
    s = sum(col_twips)
    col_twips = [int(c * AVAIL_TWIPS / s) for c in col_twips]

    t = doc.add_table(rows=len(rows), cols=cols)
    t.autofit = False
    _set_table_layout_fixed(t, AVAIL_TWIPS)
    _set_table_grid(t, col_twips)
    if rows:
        _repeat_table_header(t.rows[0])
    for ri, row in enumerate(rows):
        _set_row_cant_split(t.rows[ri])
        for ci in range(cols):
            cell = t.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_margins(cell)
            # set width explicitly
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(col_twips[ci]))
            tcW.set(qn("w:type"), "dxa")

            txt = row[ci] if ci < len(row) else ""
            p = cell.paragraphs[0]
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if display_len(txt.strip()) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_rhythm(p, after=Pt(1), line_spacing=1.08)
            # smaller font for table cells to avoid overflow
            r = p.add_run(txt.strip())
            cell_size = Pt(9) if cols >= 4 or display_len(txt) > 48 else Pt(10)
            if ri == 0:
                set_run_font(r, size=cell_size, bold=True)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), TABLE_FILL)
                tcPr.append(shd)
            else:
                set_run_font(r, size=cell_size)
            _set_cell_border(cell)
    # spacing after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def render_official_cover(doc, cfg):
    """Render cover in the OFFICIAL competition format (no decoration)."""
    # Header (matches v1 format: two-line title)
    top = doc.add_paragraph()
    set_paragraph_rhythm(top, after=Pt(42))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rhythm(p, after=Pt(2))
    r = p.add_run("臺北市 114 年度中小學資通訊應用大賽")
    set_run_font(r, size=Pt(20), bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rhythm(p, after=Pt(0))
    r = p.add_run("智組型機器人")
    set_run_font(r, size=Pt(20), bold=False)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("創意賽作品說明書")
    set_run_font(r, size=Pt(20), bold=False)
    set_paragraph_rhythm(p, after=Pt(26))

    # 學層別 line (with checkboxes, like v1)
    p = doc.add_paragraph()
    if cfg["level"] == "國小":
        text = "學 層 別：☑國小 □國中 □高中"
    elif cfg["level"] == "國中":
        text = "學 層 別：□國小 ☑國中 □高中"
    else:
        text = "學 層 別：□國小 □國中 ☑高中"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rhythm(p, before=Pt(8), after=Pt(0), line_spacing=1.08)
    r = p.add_run(text)
    set_run_font(r, size=Pt(16), bold=False)

    # 作品名稱
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_rhythm(p, before=Pt(12), after=Pt(0), line_spacing=1.08)
    r = p.add_run(f"作品名稱：{cfg['name']}")
    set_run_font(r, size=Pt(16), bold=False)

    # 關鍵詞
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.65)
    set_paragraph_rhythm(p, before=Pt(22), after=Pt(0), line_spacing=1.08)
    r = p.add_run(f"關 鍵 詞：{cfg['keywords']}　（最多 3 個）")
    set_run_font(r, size=Pt(14), bold=False)

    # 編號
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.65)
    set_paragraph_rhythm(p, before=Pt(4), after=Pt(0), line_spacing=1.08)
    r = p.add_run("編　　號：")
    set_run_font(r, size=Pt(16), bold=False)

    # 製作說明
    p = doc.add_paragraph()
    set_paragraph_rhythm(p, before=Pt(24), after=Pt(8))
    r = p.add_run("製作說明：")
    set_run_font(r, size=Pt(12))

    notes = [
        "1. 說明書封面僅寫組別、年級別、作品名稱及關鍵詞。",
        "2. 編號由承辦單位統一編列。",
        "3. 封面由選手自行設計。",
    ]
    for n in notes:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.9)
        p.paragraph_format.first_line_indent = Cm(-0.32)
        set_paragraph_rhythm(p, before=Pt(6), after=Pt(0))
        r = p.add_run(n)
        set_run_font(r, size=Pt(12))

    # Page break
    doc.add_page_break()


def parse_and_render(doc, md, cfg, app_key):
    lines = md.splitlines()
    i = 0
    in_code = False
    code_kind = None
    code_buf = []
    cover_done = False
    skip_until_heading = False

    while i < len(lines):
        line = lines[i]

        if skip_until_heading:
            mh = HEADING_RE.match(line)
            if mh and len(mh.group(1)) <= 2:
                text = mh.group(2)
                if not any(k in text for k in ["封面", cfg["name"], "作品說明書", "草稿", "重寫"]):
                    skip_until_heading = False
                else:
                    i += 1; continue
            else:
                i += 1; continue

        m = CODE_FENCE.match(line)
        if m:
            if not in_code:
                in_code = True
                code_kind = m.group(1).strip().lower()
                code_buf = []
            else:
                in_code = False
                if code_kind == "mermaid":
                    pass
                else:
                    p = doc.add_paragraph()
                    r = p.add_run("\n".join(code_buf))
                    set_run_font(r, size=Pt(9), code=True)
                code_buf = []
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        m = IMG_RE.match(line)
        if m:
            name, kind, ref = m.group(1), m.group(2), m.group(3)
            add_image(doc, name.strip(), kind.strip(), ref.strip(), app_key, line_no=i + 1)
            i += 1; continue

        m = HEADING_RE.match(line)
        if m:
            hashes, text = m.group(1), m.group(2)
            level = len(hashes)
            if not cover_done:
                render_official_cover(doc, cfg)
                cover_done = True
                if any(k in text for k in ["封面", cfg["name"], "作品說明書", "草稿", "重寫"]):
                    skip_until_heading = True
                    i += 1; continue
            if level <= 2 and "封面" in text:
                skip_until_heading = True
                i += 1; continue
            # Markdown drafts use ## for official top-level chapters.
            doc_level = max(1, min(level - 1, 3))
            p = doc.add_heading(text.strip(), level=doc_level)
            p.paragraph_format.keep_with_next = True
            i += 1; continue

        if TABLE_HEAD.match(line):
            tbl_rows = []
            while i < len(lines) and TABLE_HEAD.match(lines[i]):
                if TABLE_SEP.match(lines[i]):
                    i += 1; continue
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                tbl_rows.append(cells)
                i += 1
            add_md_table(doc, tbl_rows)
            continue

        m = LIST_RE.match(line)
        if m:
            indent, marker, content = m.group(1), m.group(2), m.group(3)
            level = len(indent.replace("\t", "    ")) // 2
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.62 + level * 0.42)
            p.paragraph_format.first_line_indent = Cm(-0.34)
            p.paragraph_format.keep_together = True
            set_paragraph_rhythm(p, after=Pt(2), line_spacing=1.12)
            label = marker if re.match(r"\d+\.", marker) else "•"
            lr = p.add_run(label + " ")
            set_run_font(lr, size=BODY_SIZE)
            parse_inline_into_paragraph(p, content)
            i += 1; continue

        m = BLOCK_QUOTE.match(line)
        if m:
            p = doc.add_paragraph()
            r = p.add_run(m.group(1))
            set_run_font(r, size=BODY_SIZE, italic=True, color=DARK_GRAY)
            p.paragraph_format.left_indent = Cm(0.6)
            i += 1; continue

        if line.strip() == "---":
            i += 1; continue

        if not line.strip():
            i += 1; continue

        p = doc.add_paragraph()
        set_paragraph_rhythm(p)
        parse_inline_into_paragraph(p, line)
        i += 1


def build_one(app_key, cfg):
    # Use v1 as template (carries styles, fonts, table styles, page setup)
    if not cfg["template"].exists():
        raise FileNotFoundError(f"v1 template missing: {cfg['template']}")
    doc = Document(str(cfg["template"]))
    configure_styles(doc)
    wipe_body(doc)
    md = cfg["src"].read_text(encoding="utf-8")
    parse_and_render(doc, md, cfg, app_key)
    normalize_doc_fonts(doc)
    cfg["out"].parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(cfg["out"]))
    print(f"  saved: {cfg['out']}")


def main():
    print("Building v2 docx from v1 templates...")
    for app_key, cfg in APP_CONFIG.items():
        print(f"\n== {app_key} ({cfg['name']}) ==")
        build_one(app_key, cfg)
    print("\nDone.")


if __name__ == "__main__":
    main()
